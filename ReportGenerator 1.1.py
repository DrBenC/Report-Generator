"""
This script will open.par files and combine them into a single, multisheet excel file"""

import docx
import pandas as pd
import glob
import os
import openpyxl
from matplotlib import pyplot as plt
import time

#This retrieves the path of the script and searches for all .par files

abspath = os.path.abspath("ReportGenerator 1.0.py")
path = os.path.dirname(abspath)

filenames1 = glob.glob(path + "/*.par")
filenames2 = [os.path.split(y)[1] for y in filenames1]
filenames = [os.path.splitext(os.path.basename(z))[0] for z in filenames2]

#Finds and retrieves .xls file to use as FILENAME - if no .xls will simply name the files 'sample'
core_file_extpath = glob.glob(path + "/*.xls")
core_fileext = [os.path.split(h)[1] for h in core_file_extpath]
core_file = [os.path.splitext(os.path.basename(j))[0] for j in core_fileext]

core_file_s = input("Excel file not found - please manually type sample name:")

#Output excel file and word document are created

writer = pd.ExcelWriter(core_file_s+" summary.xlsx")
doc = docx.Document()
test_df = pd.read_fwf(filenames[0]+".par", delim_whitespace=True)
table_heads = ["Region", "Peak", "Species", "Peak Position / eV", "Peak FWHM / eV", "Percentage Integration"]

for file in filenames:

    title = str(file)+" region in "+str(core_file_s)
    sn_dat = pd.read_fwf(file+".dat", delim_whitespace=True)
    print("Reading "+file+".dat...")

    #This is used later to set the x limits for the graph
    energy = list(sn_dat["B.E.(eV)"])
    energy_reverse = list(reversed(energy))

    #Sometimes the residuals are output as strings, this fixes that  
    differences_raw = list(sn_dat["difference"])
    differences = ([float(diff) if type(diff)==float else float(diff[:5]) for diff in differences_raw])
        
    #loading message to show data has been taken in and processed successfully
    print("Drawing "+file+"...")

    #draws shortlist of column headings in dataframe
    heads = list(sn_dat.columns)

    #confirms no open plot and sets figure size
    plt.close()
    fig = plt.figure(figsize = (6.5, 4.5))

    #works through each peak within the file
    for entry in sn_dat.columns:

        #Sets off first subplot (data), and plots all series except B.E. and Difference, sets x limits to reverse x axis
        ax_signal = plt.subplot(4,2,(1,6), label=file+" signal plot")
        if entry != "B.E.(eV)":
            if entry != "difference":
              
                #If Peak Sum data, will be a dashed line to aid visualisation
                if entry == "Peak Sum ":
                    ax_signal.plot(energy, sn_dat[entry], label=entry, linestyle="--")
                    ax_signal.set_xlim(energy[0], energy[-1])
                else:
                    ax_signal.plot(energy, sn_dat[entry], label=entry)
                    ax_signal.set_xlim(energy[0], energy[-1])
                    print(str(energy[-1])+"left")

                #Formats first subplot
                plt.tick_params(labelleft=False, left=False)
                plt.legend()
                plt.xticks(fontsize=12)
                plt.ylabel("Counts / Arbitrary", fontsize = 12)                   
                plt.ytick_labels=False
                
                #Plots second subplot (residuals from fit) and formats
            else:
                
                ax_difference = plt.subplot(4,2,(7,8), label=entry+" difference plot")
                ax_difference.plot(energy, list(differences))
                ax_difference.set_xlim(energy[0], energy[-1])
                plt.tick_params(labelleft=False, left=False)
                plt.xticks(fontsize=12)
                plt.xlabel("Binding Energy / eV", fontsize = 12)
                plt.ylabel("Residuals", fontsize = 12)
                       
    #final layout fixing, showing and saving the figure as a .png
    fig.tight_layout()
    fig.savefig(title+".png")
    plt.close()

    #This should trim down the title to a nice, manageable label
    file_list = file.split()
    if len(file_list) < 2:
        file_new_name = file
    else:
        file_new_name = str(file_list[0]+" "+file_list[1])

    #here we add the figure to the .docx report, along with the captions
    doc.add_heading(file_new_name+" region in "+core_file_s, level =1)
    doc.add_picture(title+".png")
    doc.add_paragraph("XPS spectrum showing "+file_new_name+" region in "+core_file_s+
                      ". Deconvolution of signal is shown in the top plot, with residual difference between fitted model and sample data shown below.")
    doc.add_paragraph("Table summarising deconvoluted signals giving rise to peaks in the "+file_new_name+" region in "
                      +core_file_s+" as shown above.")
    
    title = str(file)+" region in "+str(core_file_s)
    sn = pd.read_fwf(file+".par", delim_whitespace=True)
    print("Reading "+file+".par...")

    #calculates percentage integration and confirms totals add up to 100%
    columns = (list(sn.columns))
    ints = list(sn[columns[3]])
    int_percents = []
    total_int = sum(ints)
    for integration in ints:
        int_percents.append(round(((integration/total_int)*100), 1))
    sn["Percent Integration"] = int_percents
    
    #Total Integration Percentage Sum should be 100 as a sanity check
    sn["Total Integration Percentage Sum"] = sum(int_percents)

    #writes file to excel sheet - separate sheet for each .par file
    if len(file) > 15:
        sheet_filename = file[:14]
    else:
        sheet_filename = file
    sn.to_excel(writer, sheet_name=sheet_filename)

    #begins the empty lists used to make the new dataframe for the word tables
    region_list = [file_new_name]
    species = []

    def fill_list(list_o):
        while len(sn["Peak"]) != len(list_o):
            list_o.append(" ")

    fill_list(region_list)
    fill_list(species)    
      
    #now we tidy up the data from the .par files - gets each column reasonably uniform
    def round_list(list_n, n):
        
        new_list = []
        raw = list(sn[list_n])
        for m in raw:
            new_list.append(round(m, n))
        return new_list
       
    peaks = round_list("Peak", 1)
    positions = round_list("Position", 1)
    fwhm = round_list("FWHM (eV)", 2)
    percent = round_list("Percent Integration", 1)

    #we make a "reduced dataframe" which streamlines a lot of the data                              
    sn_reduced_data = {table_heads[0]: region_list, table_heads[1]: peaks, table_heads[2]: species, table_heads[3]: positions,
                       table_heads[4]: fwhm, table_heads[5]: percent}

    sn_r = pd.DataFrame(data = sn_reduced_data, columns = table_heads)
    
    #now we add the table to the word document
    table = doc.add_table(rows = sn_r.shape[0]+1, cols = 6)
    heading_table = table.rows[0].cells

    #this fills out the headings correctly
    for cell, heading in zip(heading_table, table_heads):
        cell.text = heading

    print("Writing "+file+" to report...")
    #this fills out the table with data from the reduced dataframe    
    for i in range(sn_r.shape[0]):
        for j in range(sn_r.shape[1]):
            table.cell(i+1, j).text = str(sn_r.values[i, j])

    #we start a new page ready for the next file
    if file != filenames[len(filenames)-1]:
        doc.add_page_break()                                          
   

#finally saves both files and outputs a success message
writer.save()
doc.save(core_file_s+" report.docx")
print(core_file_s+" summary.xlsx and "+core_file_s+" report.docx have been created successfully.")
time.sleep(3)
